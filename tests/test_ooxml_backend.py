"""OoxmlBackend 修订编辑回归测试。

用 python-docx 现场构造 docx（含刻意拆散的 run，模拟真实 Word 文档里
拼写检查等造成的碎片化），验证修订标记的正确性与可回读性。
"""

import zipfile

import pytest
from docx import Document
from docx.oxml.ns import qn

from word_agent_mcp.backend import OoxmlBackend
from word_agent_mcp.revisions import RevisionError


@pytest.fixture
def contract(tmp_path):
    """样例合同：第 2 段刻意由多个 run 拼成。"""
    doc = Document()
    doc.add_paragraph("房屋租赁合同")
    p = doc.add_paragraph()
    p.add_run("出租方（甲方）：张")
    p.add_run("三，身份证号 1101")
    p.add_run("01199001010011")
    doc.add_paragraph("租金为每月人民币 8000 元，押一付三。")
    doc.add_paragraph("本合同一式两份，甲方、乙方各执一份。")
    path = tmp_path / "contract.docx"
    doc.save(str(path))
    return str(path)


def test_read_paragraphs(contract):
    paras = OoxmlBackend(contract, backup=False).get_paragraphs()
    assert paras[0]["text"] == "房屋租赁合同"
    assert paras[1]["text"] == "出租方（甲方）：张三，身份证号 110101199001010011"


def test_replace_within_single_run(contract):
    doc = OoxmlBackend(contract)
    assert doc.replace_text("8000 元", "9000 元") == 1
    saved = doc.save()

    revs = OoxmlBackend(saved, backup=False).list_revisions()
    deleted = [r for r in revs if r["type"] == "删除"]
    inserted = [r for r in revs if r["type"] == "插入"]
    assert [r["text"] for r in deleted] == ["8000 元"]
    assert [r["text"] for r in inserted] == ["9000 元"]
    # 有效文本 = 接受修订后的结果
    paras = OoxmlBackend(saved, backup=False).get_paragraphs()
    assert paras[2]["text"] == "租金为每月人民币 9000 元，押一付三。"


def test_replace_across_runs(contract):
    """匹配区间横跨 3 个 run（"张三，身份证号 1101..."被拆在多个 run 里）。"""
    doc = OoxmlBackend(contract)
    doc.replace_text("张三，身份证号 110101199001010011", "李四，身份证号 110101198802020022")
    saved = doc.save()
    paras = OoxmlBackend(saved, backup=False).get_paragraphs()
    assert paras[1]["text"] == "出租方（甲方）：李四，身份证号 110101198802020022"
    revs = OoxmlBackend(saved, backup=False).list_revisions()
    assert "".join(r["text"] for r in revs if r["type"] == "删除") == (
        "张三，身份证号 110101199001010011"
    )


def test_new_text_containing_old_not_rematched(contract):
    """新文本包含旧文本时不得把刚插入的内容再次替换。"""
    doc = OoxmlBackend(contract)
    n = doc.replace_text("甲方", "甲方（出租人）", replace_all=True)
    assert n == 2  # 第 1 段 1 处 + 末段 1 处
    saved = doc.save()
    paras = OoxmlBackend(saved, backup=False).get_paragraphs()
    assert "出租方（甲方（出租人））" in paras[1]["text"]
    assert "甲方（出租人）、乙方" in paras[3]["text"]
    # 恰好 2 删 2 插，没有嵌套多余修订
    revs = OoxmlBackend(saved, backup=False).list_revisions()
    assert len([r for r in revs if r["type"] == "删除"]) == 2
    assert len([r for r in revs if r["type"] == "插入"]) == 2


def test_ambiguous_match_rejected(contract):
    doc = OoxmlBackend(contract)
    with pytest.raises(RevisionError, match="出现 2 次"):
        doc.replace_text("甲方", "承租方")


def test_not_found_rejected(contract):
    with pytest.raises(RevisionError, match="未找到"):
        OoxmlBackend(contract).replace_text("不存在的文字", "x")


def test_pure_deletion(contract):
    doc = OoxmlBackend(contract)
    doc.replace_text("，押一付三", "")
    saved = doc.save()
    paras = OoxmlBackend(saved, backup=False).get_paragraphs()
    assert paras[2]["text"] == "租金为每月人民币 8000 元。"


def test_insert_and_delete_paragraph(contract):
    doc = OoxmlBackend(contract)
    doc.insert_paragraph_after(2, "乙方逾期支付租金超过十五日的，甲方有权解除本合同。")
    doc.delete_paragraph(3 + 1)  # 原末段（插入后顺延为第 4 段）
    saved = doc.save()

    revs = OoxmlBackend(saved, backup=False).list_revisions()
    types = {r["type"] for r in revs}
    assert "插入段落标记" in types and "删除段落标记" in types
    ins_text = "".join(r["text"] for r in revs if r["type"] == "插入")
    assert "解除本合同" in ins_text
    del_text = "".join(r["text"] for r in revs if r["type"] == "删除")
    assert "一式两份" in del_text


def test_backup_created(contract, tmp_path):
    OoxmlBackend(contract).replace_text("8000", "8500")
    assert (tmp_path / "contract.docx.bak").exists()


def test_output_is_valid_docx(contract):
    """产出文件必须能被 python-docx 重新打开，且 XML 里确有 w:ins/w:del。"""
    doc = OoxmlBackend(contract)
    doc.replace_text("8000 元", "9000 元")
    saved = doc.save()
    Document(saved)  # 不抛异常即结构合法
    with zipfile.ZipFile(saved) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    assert "<w:ins " in xml and "<w:del " in xml and "w:delText" in xml


def test_revision_ids_unique(contract):
    doc = OoxmlBackend(contract)
    doc.replace_text("8000 元", "9000 元")
    doc.replace_text("租赁合同", "租赁协议")
    saved = doc.save()
    body = Document(saved).element.body
    ids = [
        el.get(qn("w:id"))
        for el in body.iter(qn("w:ins"), qn("w:del"))
    ]
    assert len(ids) == len(set(ids))


def test_relative_path_rejected():
    """MCP server 工作目录 ≠ 会话目录，相对路径必须被拒绝并引导 agent 纠正。"""
    from word_agent_mcp.server import read_document

    result = read_document("演示合同.docx")
    assert "错误" in result and "绝对路径" in result

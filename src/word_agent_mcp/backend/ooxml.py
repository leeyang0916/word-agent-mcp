"""OOXML 文件后端：python-docx + lxml 直接读写 .docx，修订以 w:ins/w:del 写入。

跨平台（Mac/Windows/Linux 行为一致）。要求目标文档未被 Word 以写模式占用
（Windows 上 Word 会锁定打开中的文件），推荐工作流：改副本 → Word 审阅。
"""

from __future__ import annotations

import shutil
from datetime import datetime, timezone
from pathlib import Path

from docx import Document

from .. import revisions as rev
from .base import WordBackend


class OoxmlBackend(WordBackend):
    is_live = False

    def __init__(self, path: str, *, author: str = "Word Agent", backup: bool = True):
        self.path = Path(path)
        if not self.path.exists():
            raise FileNotFoundError(f"文档不存在: {path}")
        if backup:
            bak = self.path.with_suffix(self.path.suffix + ".bak")
            if not bak.exists():
                shutil.copy2(self.path, bak)
        self.doc = Document(str(self.path))
        self._body = self.doc.element.body
        date = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
        self.ctx = rev.RevisionContext.for_document(self._body, author=author, date=date)

    # ---- 读取 ----

    def get_paragraphs(self) -> list[dict]:
        return [
            {
                "index": i,
                "style": p.style.name if p.style else None,
                "text": rev.paragraph_effective_text(p._p),
            }
            for i, p in enumerate(self.doc.paragraphs)
        ]

    def list_revisions(self) -> list[dict]:
        return rev.list_revisions(self._body)

    # ---- 修订式编辑 ----

    def replace_text(self, old: str, new: str, *, replace_all: bool = False) -> int:
        if not old:
            raise rev.RevisionError("old 不能为空")
        hits = [(p, rev.count_occurrences(p._p, old)) for p in self.doc.paragraphs]
        total = sum(n for _, n in hits)
        if total == 0:
            raise rev.RevisionError(f"全文未找到文本: {old!r}")
        if total > 1 and not replace_all:
            raise rev.RevisionError(
                f"文本 {old!r} 出现 {total} 次。请提供更长的唯一上下文，"
                f"或明确指定 replace_all=True 全部替换"
            )
        count = 0
        for p, n in hits:
            if n:
                count += rev.tracked_replace_in_paragraph(
                    p._p, old, new, self.ctx, all_occurrences=True
                )
        return count

    def insert_paragraph_after(self, index: int, text: str, *, style: str | None = None) -> None:
        p = self._paragraph_at(index)
        rev.tracked_new_paragraph(p._p, text, self.ctx, style=style)

    def delete_paragraph(self, index: int) -> None:
        p = self._paragraph_at(index)
        rev.tracked_delete_paragraph(p._p, self.ctx)

    # ---- 持久化 ----

    def save(self, path: str | None = None) -> str:
        target = str(path or self.path)
        self.doc.save(target)
        return target

    # ---- 内部 ----

    def _paragraph_at(self, index: int):
        paras = self.doc.paragraphs
        if not 0 <= index < len(paras):
            raise rev.RevisionError(f"段落序号 {index} 超出范围（共 {len(paras)} 段，从 0 计数）")
        return paras[index]

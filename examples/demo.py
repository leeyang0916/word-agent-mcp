"""最小演示：生成一份样例合同 → agent 式修订编辑 → 输出待审阅的 docx。

运行: uv run python examples/demo.py
产出: examples/output/演示合同.docx（原稿）与 演示合同_已修订.docx
用 Word / Pages 打开"已修订"版即可看到标准红线修订。
"""

from pathlib import Path

from docx import Document

from word_agent_mcp.backend import OoxmlBackend

OUT = Path(__file__).parent / "output"
OUT.mkdir(exist_ok=True)
original = OUT / "演示合同.docx"
revised = OUT / "演示合同_已修订.docx"

# 1. 造一份原稿
doc = Document()
doc.add_heading("房屋租赁合同", level=1)
doc.add_paragraph("出租方（甲方）：张三")
doc.add_paragraph("承租方（乙方）：李四")
doc.add_paragraph("第一条 租金为每月人民币 8000 元，押一付三，先付后住。")
doc.add_paragraph("第二条 租赁期限自 2026 年 8 月 1 日起至 2027 年 7 月 31 日止。")
doc.add_paragraph("第三条 本合同一式两份，甲乙双方各执一份，自签字之日起生效。")
doc.save(str(original))
print(f"原稿: {original}")

# 2. 模拟 agent 的修订式编辑（实际产品中这些调用来自 LLM 的工具调用）
import shutil

shutil.copy2(original, revised)
agent_doc = OoxmlBackend(str(revised), author="AI 律师助手", backup=False)
agent_doc.replace_text("8000 元", "9000 元")
agent_doc.replace_text(
    "押一付三，先付后住。",
    "押一付三，先付后住。甲方收取押金后应向乙方出具收据。",
)
agent_doc.insert_paragraph_after(
    5,
    "第四条 乙方逾期支付租金超过十五日的，甲方有权解除本合同并要求乙方承担违约责任。",
)
agent_doc.save()

# 3. 汇报产生的修订
print(f"修订稿: {revised}\n")
print("本次产生的修订：")
for r in agent_doc.list_revisions():
    text = f"“{r['text']}”" if r["text"] else ""
    print(f"  [{r['type']}] by {r['author']} {text}")
print("\n请用 Word 打开修订稿，在「审阅」里逐条接受/拒绝。")
